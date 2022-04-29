#This file has code committed by Luis Carrillo, David Garcia, Mason Lanham, and Chris Mendoza. Specific sections are commented to show who they were committed by.

#The following section was committed by Luis Carrillo and Mason Lanham
import docx
import docx2txt as d2t
import os
from TextualElement import *
from GraphicalElement import *
from Table import *
from Cell import *
from Row import *
from Column import *

#The following section was committed by Luis Carrillo and Mason Lanham
class Parser():
    def __init__(self, docName):
        self.doc = docx.Document(docName)
        self.all_paras = self.doc.paragraphs
        self.all_tables = self.doc.tables
        sections = self.doc.sections
        self.firstPageHeader = sections[0].first_page_header 
        self.firstPageFooter = sections[0].first_page_footer
        self.header = sections[0].header
        self.footer = sections[0].footer


# Method definitions

# Text Methods -------------------------------------------------------------

#The following section was committed by Luis Carrillo
    def findHBF(self, para):
        headerParagraphs = self.header.paragraphs + self.firstPageHeader.paragraphs
        for headerPara in headerParagraphs:
            if para.text == headerPara.text:
                return 1
        footerParagraphs = self.footer.paragraphs + self.firstPageFooter.paragraphs
        for footerPara in footerParagraphs:
            if para.text == footerPara.text:
                return 2
        return 0

#The following section was commmitted by Luis Carrillo
    def findRuns(self, para):
        runsArray = []
        for run in para.runs:
            runsArray.append(run.text)
        return runsArray

#Run Methods -----------------------------------------------------------

#The following section was committed by Luis Carrillo
    def findText(self, run):
        runText = run.text
        return runText

#The following section was committed by Luis Carrillo
    def findBold(self, run):
        if run.bold:
           return bool(True)
        else:
           return bool(False)

#The following section was committed by Mason Lanham
    def findFont(self, run):
        if not(run.font.name is None):
            return run.font.name
        elif not(run.style.font.name is None):
            return run.style.font.name
        else:
            return "Unknown"

#The following section was committed by Luis Carrillo
    def findItalic(self, run):
        if run.italic:
            return bool(True)
        else:
            return bool(False)
        
#The following section was committed by Mason Lanham
    def findStyle(self, run):
        if not(run.style.name is None):
            return run.style.name
        else:
            return "Unknown"

#The following section was committed by Luis Carrillo
    def findUnderline(self, run):
        if run.underline:
            return bool(True)
        else:
            return bool(False)

#Table Methods ---------------------------------------------------------

#The following section was committed by Luis Carrillo
    def findColumnsList(self, table):
        columnList = []
        columnNum = 0
        for column in table.columns:
            columnList.append(Column(self.findTableName(table), "", self.findNumOfCells(column)))
            columnName = column.cells[0].text.replace("\n", " ")
            columnList[columnNum].setColumnName(columnName)
            columnNum += 1
        return columnList

#The following section was committed by Luis Carrillo
    def findRowsList(self, table):
        rowList = []
        rowNum = 0
        for row in table.rows:
            rowList.append(Row(self.findTableName(table), "", self.findNumOfCells(row)))
            rowName = row.cells[0].text.replace("\n", " ")
            rowList[rowNum].setRowName(rowName)
            rowNum += 1
        return rowList 

#The following section was committed by Mason Lanham and Chris Mendoza
    def findCellsList(self, table): #Method returns a list of Cell objects that make up the table passed to it.
        cellsList = []
        tableName = self.findTableName(table)
        i = 0
        iteration = 0
        for row in table.rows:
            j = 0
            for cell in row.cells:
                cellsList.append(Cell([], tableName, i, j))
                jteration = 0
                for para in cell.paragraphs:
                    if(len(para.runs) > 0 and para.runs[0].text != ""):
                        cellsList[iteration].appendText(TextualElement(jteration, self.findHBF(para), para.style.name, []))
                        numRuns = 0
                        for run in para.runs:
                            if numRuns > 0 :
                                run0 = cellsList[iteration].getTextualElements()[jteration].getRun(numRuns - 1)
                                if (run0.getFont() == self.findFont(run) and run0.getStyle() == self.findStyle(run) and run0.getBold() == self.findBold(run) and run0.getItalic() ==  self.findItalic(run) and run0.getUnderline() == self.findUnderline(run)):
                                    run0.appendText(self.findText(run))
                                else:
                                    cellsList[iteration].getTextualElements()[jteration].appendRun(Run(self.findText(run), self.findFont(run), self.findStyle(run), self.findBold(run), self.findItalic(run), self.findUnderline(run)))
                                    numRuns += 1
                            else:
                                cellsList[iteration].getTextualElements()[jteration].appendRun(Run(self.findText(run), self.findFont(run), self.findStyle(run), self.findBold(run), self.findItalic(run), self.findUnderline(run)))
                                numRuns += 1
                        jteration += 1
                j += 1
                iteration += 1
            i += 1
        return cellsList

#The following section was committed by Luis Carrillo
    def findNumOfColumns(self, table):
        numOfColumns = 0
        for column in table.columns:
            numOfColumns += 1
        return numOfColumns

#The following section was committed by Luis Carrillo
    def findNumOfRows(self, table):
        numOfRows = 0
        for row in table.rows:
            numOfRows += 1
        return numOfRows

#The following section was committed by Luis Carrillo    
    def findTableName(self, table):
        tableName = ""
        if table.rows[0].cells[0].text and table.rows[0].cells[0].text.strip():
            tableName += table.rows[0].cells[0].text
        else:
            tableName += table.rows[0].cells[1].text
        return tableName

#Row Methods -----------------------------------------------------------

#The following section was committed by Luis Carrillo
    def findRowString(self, row):
        rowString = ""
        for cell in row.cells:
            for para in cell.paragraphs:
                rowString += para.text
                rowString += " "
        return rowString

#The following section was committed by Luis Carrillo
    def findTable(self, array):
        return array.table

#The following section was committed by Luis Carrillo    
    def findNumOfCells(self, array):
        cellCount = 0
        for cell in array.cells:
            cellCount += 1
        return cellCount


#Cell Methods -----------------------------------------------------------------

#The following section was committed by Luis Carrillo
    def c_findText(self, cell):
        cellText = ""
        for para in cell.paras:
            cellText += para.text
        if cellText == "":
            cellText = "_"
        return cellText
    
#The following section was committed by Mason Lanham and Chris Mendoza    
    def tablesParse(self): #Method returns a list of table objects containing all tables in the Word document
        tableList = []
        iteration = 0
        for table in self.all_tables:
            tableList.append(Table([], [], [], 0, 0, iteration, self.findTableName(table)))
            tableList[iteration].setColumns(self.findColumnsList(table))
            tableList[iteration].setRows(self.findRowsList(table))
            tableList[iteration].setCells(self.findCellsList(table))
            tableList[iteration].setNumberOfColumns(tableList[iteration].getRows()[0].getNumberOfCells())
            tableList[iteration].setNumberOfRows(tableList[iteration].getColumns()[0].getNumberOfCells())
            iteration += 1
        return tableList

#The following section was committed by Luis Carillo and Mason Lanham
    def paragraphsParse(self): #Method returns a list of TextualElement objects containing all paragraphs not contained within tables in the Word document
        textualElementList = []
        iteration = 0
        for para in self.all_paras:
            if(len(para.runs) > 0 and para.runs[0].text != ""):
                textualElementList.append(TextualElement(iteration, self.findHBF(para), para.style.name, []))
                numRuns = 0
                for run in para.runs:
                    if numRuns > 0 :
                        run0 = textualElementList[iteration].getRun(numRuns - 1)
                        if (run0.getFont() == self.findFont(run) and run0.getStyle() == self.findStyle(run) and run0.getBold() == self.findBold(run) and run0.getItalic() ==  self.findItalic(run) and run0.getUnderline() == self.findUnderline(run)):
                            run0.appendText(self.findText(run))
                        else:
                            textualElementList[iteration].appendRun(Run(self.findText(run), self.findFont(run), self.findStyle(run), self.findBold(run), self.findItalic(run), self.findUnderline(run)))
                            numRuns += 1
                    else:
                        textualElementList[iteration].appendRun(Run(self.findText(run), self.findFont(run), self.findStyle(run), self.findBold(run), self.findItalic(run), self.findUnderline(run)))
                        numRuns += 1
                iteration += 1
        return textualElementList

#The following section was committed by David Garcia and Mason Lanham
    def graphicsParse(self, inputFilePath, outputFileLocation): #Method returns a list of GraphicalElement objects containing all graphics in the Word document
        self.extract_images(inputFilePath, outputFileLocation)
        GraphicsList = []
        iteration = 0
        for file in os.listdir(outputFileLocation):
            if file.endswith('.jpeg') or file.endswith('.png') or file.endswith('.gif') or file.endswith('.bmp'):
                GraphicsList.append(GraphicalElement(iteration, outputFileLocation + file))
                iteration += 1
        return GraphicsList

#The following section was committed by David Garcia 
    def extract_images(self, path_to_file, images_folder, get_text=False):
        text = d2t.process(path_to_file, images_folder)
        if(get_text):
            return text
    
