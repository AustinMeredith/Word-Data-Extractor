#Purpose is to see if I could extract the text from a document I manually edited, save in a .xml
import lxml
from docx import Document

doc = Document('Test.docx')
new = Document()
for paragraph in doc.paragraphs:
    new.add_paragraph(paragraph.text)
new.save('New.xml')
