#Purpose is to see if I could extract the text from a document I manually edited, save in a .xml and then retrieve that text from the .xml and display in a new word doc
import lxml
from docx import Document

doc = Document('New.xml')
new = Document()
for paragraph in doc.paragraphs:
    new.add_paragraph(paragraph.text)
new.save('New.docx')
